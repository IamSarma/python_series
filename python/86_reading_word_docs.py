import docx
import read_docx

# doc = docx.Document("demo.docx")

# # Getting the length of the paragraphs in the document
# print(len(doc.paragraphs))

# # Getting text using index of the paragraphs object
# print(doc.paragraphs[0].text)
# print(doc.paragraphs[1].text)

# # Getting number of runs in the paragraphs object
# print(len(doc.paragraphs[1].runs))

# # Getting text from the runs object using index and text
# print(doc.paragraphs[1].runs[0].text)
# print(doc.paragraphs[1].runs[1].text)
# print(doc.paragraphs[1].runs[2].text)
# print(doc.paragraphs[1].runs[3].text)
# print(doc.paragraphs[1].runs[4].text)


# Extract and display full text from the given word document
# print(read_docx.getText("demo.docx"))


# # Applying style(s)
# print(doc.paragraphs[0].text)
# print(doc.paragraphs[0].style)
# doc.paragraphs[0].style = "Normal"

# print(doc.paragraphs[1].text)
# print(doc.paragraphs[1].runs[0].text,
#       doc.paragraphs[1].runs[1].text,
#       doc.paragraphs[1].runs[2].text,
#       doc.paragraphs[1].runs[3].text,
#       doc.paragraphs[1].runs[4].text)
# doc.paragraphs[1].runs[0].QuoteChar = True
# doc.paragraphs[1].runs[2].underline = True
# doc.paragraphs[1].runs[4].underline = True
# doc.save("restyled.docx")


# # Writing word document(s)
# doc = docx.Document()
# doc.add_paragraph("Python is awesome!!!", "Title")
# doc.save("made_with_python.docx")

# para_obj1 = doc.add_paragraph("This is a second paragraph.")
# para_obj2 = doc.add_paragraph("This is yet another paragraph.")
# para_obj1.add_run(" This text is being added to the second paragraph.")
# doc.save("multiple_paragraphs.docx")


# # Adding heading(s)
# doc = docx.Document()
# doc.add_heading("Header 0", 0)
# doc.add_heading("Header 0", 1)
# doc.add_heading("Header 0", 2)
# doc.add_heading("Header 0", 3)
# doc.add_heading("Header 0", 4)
# doc.save("headings.docx")


# Adding page break
doc = docx.Document()
doc.add_paragraph("This paragraph is on the first page.")
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph("This paragraph is on the second page.")
doc.save("two_pages.docx")
